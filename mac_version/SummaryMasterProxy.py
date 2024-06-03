import os
from openai import OpenAI
import PyPDF2
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import tkinter as tk
from tkinter import ttk, filedialog, messagebox

def create_gui():
    root = tk.Tk()
    root.title("PDF 文章批量总结生成器")

    # 创建输入框和按钮
    folder_label = tk.Label(root, text="PDF 文件夹路径:")
    folder_label.grid(row=0, column=0, padx=10, pady=10)
    folder_entry = tk.Entry(root)
    folder_entry.grid(row=0, column=1, padx=10, pady=10)
    folder_button = tk.Button(root, text="选择文件夹", command=lambda: select_folder(folder_entry))
    folder_button.grid(row=0, column=2, padx=10, pady=10)

    markdown_label = tk.Label(root, text="Markdown 输出文件夹:")
    markdown_label.grid(row=1, column=0, padx=10, pady=10)
    markdown_entry = tk.Entry(root)
    markdown_entry.grid(row=1, column=1, padx=10, pady=10)
    markdown_button = tk.Button(root, text="选择文件夹", command=lambda: select_folder(markdown_entry))
    markdown_button.grid(row=1, column=2, padx=10, pady=10)

    word_label = tk.Label(root, text="Word 输出文件夹:")
    word_label.grid(row=2, column=0, padx=10, pady=10)
    word_entry = tk.Entry(root)
    word_entry.grid(row=2, column=1, padx=10, pady=10)
    word_button = tk.Button(root, text="选择文件夹", command=lambda: select_folder(word_entry))
    word_button.grid(row=2, column=2, padx=10, pady=10)

    api_label = tk.Label(root, text="OpenAI API 密钥:")
    api_label.grid(row=3, column=0, padx=10, pady=10)
    api_entry = tk.Entry(root)
    api_entry.grid(row=3, column=1, padx=10, pady=10)

    base_label = tk.Label(root, text="OpenAI API 基础 URL:")
    base_label.grid(row=4, column=0, padx=10, pady=10)
    base_entry = tk.Entry(root)
    base_entry.insert(0, "https://api.moonshot.cn/v1")
    base_entry.grid(row=4, column=1, padx=10, pady=10)
    # 添加代理输入框
    http_proxy_label = tk.Label(root, text="http 代理地址:")
    http_proxy_label.grid(row=6, column=0, padx=10, pady=10)
    http_proxy_entry = tk.Entry(root)
    http_proxy_entry.insert(0, "http://127.0.0.1:7890")
    http_proxy_entry.grid(row=6, column=1, padx=10, pady=10)
    http_proxy_entry.config(state="disabled")

    https_proxy_label = tk.Label(root, text="https 代理地址:")
    https_proxy_label.grid(row=6, column=2, padx=10, pady=10)
    https_proxy_entry = tk.Entry(root)
    https_proxy_entry.insert(0, "http://127.0.0.1:7890")
    https_proxy_entry.grid(row=6, column=3, padx=10, pady=10)
    https_proxy_entry.config(state="disabled")

    all_proxy_label = tk.Label(root, text="socks5 代理地址:")
    all_proxy_label.grid(row=7, column=0, padx=10, pady=10)
    all_proxy_entry = tk.Entry(root)
    all_proxy_entry.insert(0, "socks5://127.0.0.1:7890")
    all_proxy_entry.grid(row=7, column=1, padx=10, pady=10)
    all_proxy_entry.config(state="disabled")
    
    # 添加是否使用代理的选项
    proxy_var = tk.BooleanVar()
    proxy_checkbox = tk.Checkbutton(root, text="使用代理", variable=proxy_var, command= lambda:toggle_proxy_entry(http_proxy_entry, https_proxy_entry,all_proxy_entry,proxy_var)) 
    proxy_checkbox.deselect()
    proxy_checkbox.grid(row=5, column=0, padx=10, pady=10)


    # 添加生成 Markdown 或 Word 文件的选项
    output_format_var = tk.StringVar()
    output_format_var.set("markdown")
    markdown_radio = tk.Radiobutton(root, text="生成 Markdown 文件", variable=output_format_var, value="markdown")
    word_radio = tk.Radiobutton(root, text="生成 Word 文件", variable=output_format_var, value="word")
    markdown_radio.grid(row=9, column=0, padx=10, pady=10)
    word_radio.grid(row=9, column=1, padx=10, pady=10)

    progress_label = tk.Label(root, text="处理进度:")
    progress_label.grid(row=10, column=0, padx=10, pady=10)
    progress_bar = ttk.Progressbar(root, mode='determinate', length=200)
    progress_bar.grid(row=10, column=1, padx=10, pady=10)

    start_button = tk.Button(root, text="开始生成", command=lambda: start_processing(
        folder_entry.get(),
        markdown_entry.get(),
        word_entry.get(),
        api_entry.get(),
        base_entry.get(),
        proxy_var.get(),
        http_proxy_entry.get(),
        https_proxy_entry.get(),
        all_proxy_entry.get(),
        output_format_var.get(),
        progress_bar,
        root
    ))
    start_button.grid(row=11, column=1, padx=10, pady=10)

    root.mainloop()

def toggle_proxy_entry(http_proxy_entry,https_proxy_entry, all_proxy_entry, proxy_var):
    if proxy_var.get():
        http_proxy_entry.config(state="normal")
        https_proxy_entry.config(state="normal")
        all_proxy_entry.config(state="normal")
    else:
        http_proxy_entry.config(state="disabled")
        https_proxy_entry.config(state="disabled")
        all_proxy_entry.config(state="disabled")

def select_folder(entry):
    folder_path = filedialog.askdirectory()
    entry.delete(0, tk.END)
    entry.insert(0, folder_path)

def start_processing(folder_path, 
                     markdown_output_folder, 
                     word_output_folder, api_key, base_url, 
                     use_proxy, http_proxy_entry, https_proxy_entry,all_proxy_entry,output_format, progress_bar, root):
    # 设置代理
    if use_proxy:
        os.environ['http_proxy'] = http_proxy_entry
        os.environ['https_proxy'] = https_proxy_entry
        os.environ['all_proxy'] = all_proxy_entry
    else:
        os.environ.pop('http_proxy', None)
        os.environ.pop('https_proxy', None)
        os.environ.pop('all_proxy', None)

    client = OpenAI(
        api_key=api_key,
        base_url=base_url,
    )

    total_files = len([f for f in os.listdir(folder_path) if f.endswith(".pdf")])
    processed_files = 0

    for filename in os.listdir(folder_path):
        if filename.endswith(".pdf"):
            file_path = os.path.join(folder_path, filename)
            print(f"正在处理文件：{file_path}")
            try:
                text = extract_text_from_pdf(file_path, max_pages=15)
                summary = summarize_text(text, filename, client)
                if output_format == "markdown":
                    save_summaries_to_markdown(filename, summary, markdown_output_folder)
                elif output_format == "word":
                    save_summaries_to_word(filename, summary, word_output_folder)
            except Exception as e:
                print(f"处理文件 {filename} 时出现错误: {e}")
            processed_files += 1
            progress_bar['value'] = (processed_files / total_files) * 100
            root.update_idletasks()

    messagebox.showinfo("处理完成", "所有 PDF 文件已经处理完毕!")

def extract_text_from_pdf(file_path, max_pages=20):
    text = ""
    with open(file_path, 'rb') as file:
        pdf_reader = PyPDF2.PdfReader(file)
        for page_num in range(min(len(pdf_reader.pages), max_pages)):
            page = pdf_reader.pages[page_num]
            text += page.extract_text()
    return text

def summarize_text(text, filename, client):
    prompt = f"请对{filename}中的论文内容进行总结,提取出关键的内容,首先在文章开头列出**文章信息**: 1.作者\n 2.文章发表时间\n,3.文章发表机构\n; 其次提取出文章正文关键内容的总结,包括但不限于:\n\n 1. 研究背景和目的\n2. 主要的方法和实验设计\n3. 主要发现和结果\n 4.核心公式(使用latex语法,渲染符号使用美元符$,并对公式进行解释)5. 结论和意义\n6.文章当前有何缺陷及以后的工作\n\n总结: 正文总结内容不少于800字,尽量地提升信息含量.\n\n 此外,请你将文章中你认为最关键和相关的3条引用的文献按照如下格式分条整理,每个文献占据一行: 1. 文献时间 2.文献标题(使用原文英文) 3. 文献主题(需根据名称总结) 4.文献期刊简称 5. 在文中出现的次数 \n\n"
    response = client.chat.completions.create(
        model="moonshot-v1-32k",
        messages=[
            {"role": "system", "content": "你是一个高效的学术总结助手。你的任务是通读一篇学术论文，并用中文给出高信息含量的总结。 公式：对于行内公式，使用单个的美元符号`$`符号进行渲染,对于行间公式,使用两对美元符号`$$`加在公式两边进行渲染，比如: 行间公式：$$ y = ax + b $$ ,除了文章标题使用一级标题, 其余板块请你只使用加粗而非多级标题来区分不同板块"},
            {"role": "system", "content": text},
            {"role": "user", "content": prompt},
        ],
        temperature=0.1
    )
    summary = response.choices[0].message.content.strip()
    return summary

def save_summaries_to_markdown(filename, summary, output_folder):
    if not os.path.exists(output_folder):
        os.makedirs(output_folder)
    markdown_filename = f"{os.path.splitext(filename)[0]}.md"
    markdown_path = os.path.join(output_folder, markdown_filename)
    if os.path.exists(markdown_path):
        if messagebox.askyesno("文件已存在", f"文件 {markdown_filename} 已存在,是否覆盖?"):
            with open(markdown_path, 'w', encoding='utf-8') as file:
                file.write(f"#{filename.split('.')[0]}\n\n## 总结:\n{summary}\n")
    else:
        with open(markdown_path, 'w', encoding='utf-8') as file:
            file.write(f"#{filename.split('.')[0]}\n\n## 总结:\n{summary}\n")

def save_summaries_to_word(filename, summary, output_folder):
    if not os.path.exists(output_folder):
        os.makedirs(output_folder)
    word_filename = f"{os.path.splitext(filename)[0]}.docx"
    word_path = os.path.join(output_folder, word_filename)
    if os.path.exists(word_path):
        if messagebox.askyesno("文件已存在", f"文件 {word_filename} 已存在,是否覆盖?"):
            document = Document()
            set_word_document_style(document)
            document.add_heading(f'文件：{filename}', level=1)
            document.add_heading('总结：', level=2)
            paragraph = document.add_paragraph(summary)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            document.save(word_path)
    else:
        document = Document()
        set_word_document_style(document)
        document.add_heading(f'文件：{filename}', level=1)
        document.add_heading('总结：', level=2)
        paragraph = document.add_paragraph(summary)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        document.save(word_path)

def set_word_document_style(document):
    # 设置标题格式
    title_style = document.styles['Heading 1']
    title_style.font.name = 'Times New Roman'
    title_style.font.size = Pt(14)
    title_style.font.bold = True

    # 设置正文格式
    body_style = document.styles['Normal']
    body_style.font.name = 'Times New Roman'
    body_style.font.size = Pt(12)
    body_style.font.bold = False
create_gui()