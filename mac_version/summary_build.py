# -*- coding: utf-8 -*-

import os
from openai import OpenAI
import PyPDF2
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

# 获取用户输入
folder_path = input("请输入包含 PDF 文件的文件夹路径: ")
markdown_output_folder = input("请输入 Markdown 输出文件夹路径: ")
api_key = input("请输入 OpenAI API 密钥: ")
base_url = input("请输入 OpenAI API 基础 URL: ")

# 设置 OpenAI 客户端
client = OpenAI(api_key=api_key, base_url=base_url,)

# 定义解析 PDF 文件的函数
def extract_text_from_pdf(file_path, max_pages=20):
    text = ""
    with open(file_path, "rb") as file:
        pdf_reader = PyPDF2.PdfReader(file)
        for page_num in range(min(len(pdf_reader.pages), max_pages)):
            page = pdf_reader.pages[page_num]
            text += page.extract_text()
    return text


# 定义生成内容总结的函数
def summarize_text(text, filename):
    prompt = f"请对{filename}中的论文内容进行总结,提取出关键的内容,首先在文章开头列出**文章信息**: 1.作者\n 2.文章发表时间\n,3.文章发表机构\n; 其次提取出文章正文关键内容的总结,包括但不限于:\n\n 1. 研究背景和目的\n2. 主要的方法和实验设计\n3. 主要发现和结果\n 4.核心公式(使用latex语法,渲染符号使用美元符$,并对公式进行解释)5. 结论和意义\n6.文章当前有何缺陷及以后的工作\n\n总结: 正文总结内容不少于800字,尽量地提升信息含量.\n\n 此外,请你将文章中你认为最关键和相关的3条引用的文献按照如下格式分条整理,每个文献占据一行: 1. 文献时间 2.文献标题(使用原文英文) 3. 文献主题(需根据名称总结) 4.文献期刊简称 5. 在文中出现的次数 \n\n"
    response = client.chat.completions.create(
        model="moonshot-v1-32k",
        messages=[
            {
                "role": "system",
                "content": "你是一个高效的学术总结助手。你的任务是通读一篇学术论文，并用中文给出高信息含量的总结。 公式：对于行内公式，使用单个的美元符号`$`符号进行渲染,对于行间公式,使用两对美元符号`$$`加在公式两边进行渲染，比如: 行间公式：$$ y = ax + b $$ ,除了文章标题使用一级标题, 其余板块请你只使用加粗而非多级标题来区分不同板块",
            },
            {"role": "system", "content": text},
            {"role": "user", "content": prompt},
        ],
        temperature=0.1,
    )
    summary = response.choices[0].message.content.strip()
    return summary


# 保存总结为 Markdown 文件
def save_summaries_to_markdown(filename, summary, output_folder):
    if not os.path.exists(output_folder):
        os.makedirs(output_folder)
    markdown_filename = f"{os.path.splitext(filename)[0]}.md"
    markdown_path = os.path.join(output_folder, markdown_filename)
    with open(markdown_path, "w", encoding="utf-8") as file:
        file.write(f"#{filename.split('.')[0]}\n\n## 总结:\n{summary}\n")


def save_summaries_to_word(filename, summary, output_folder):
    if not os.path.exists(output_folder):
        os.makedirs(output_folder)
    word_filename = f"{os.path.splitext(filename)[0]}.docx"
    word_path = os.path.join(output_folder, word_filename)
    document = Document()

    # 设置标题格式
    title_style = document.styles["Heading 1"]
    title_style.font.name = "Times New Roman"
    title_style.font.size = Pt(14)
    title_style.font.bold = True

    # 设置正文格式
    body_style = document.styles["Normal"]
    body_style.font.name = "Times New Roman"
    body_style.font.size = Pt(12)
    body_style.font.bold = False

    document.add_heading(f"文件：{filename}", level=1)
    document.add_heading("总结：", level=2)

    # 添加正文段落
    paragraph = document.add_paragraph(summary)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    document.save(word_path)


# 逐个文件生成总结
for filename in os.listdir(folder_path):
    if filename.endswith(".pdf"):
        file_path = os.path.join(folder_path, filename)
        print(f"正在处理文件：{file_path}")
        text = extract_text_from_pdf(file_path, max_pages=15)
        summary = summarize_text(text, filename)
        save_summaries_to_markdown(filename, summary, markdown_output_folder)
        # save_summaries_to_word(filename, summary, word_output_folder)
