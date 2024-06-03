import os
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import markdown

def convert_markdown_to_word(markdown_file, output_folder):
    """
    将 Markdown 文件转换为 Word 文件
    """
    # 读取 Markdown 文件内容
    with open(markdown_file, 'r', encoding='utf-8') as file:
        markdown_content = file.read()

    # 将 Markdown 转换为 HTML
    html_content = markdown.markdown(markdown_content)

    # 创建 Word 文档
    document = Document()

    # 设置标题格式
    title_style = document.styles['Heading 1']
    title_style.font.name = 'Times New Roman'
    title_style.font.size = Pt(14)
    title_style.font.bold = True

    # 设置正文格式
    body_style = document.styles['Normal']
    body_style.font.name = 'Times New Roman'
    body_style.font.size = Pt(12)
    body_style.font.bold = False

    # 添加标题
    filename = os.path.splitext(os.path.basename(markdown_file))[0]
    document.add_heading(f'文件：{filename}', level=1)

    # 添加正文
    paragraph = document.add_paragraph('')
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.add_run(html_content)

    # 保存 Word 文件
    word_filename = f"{filename}.docx"
    word_path = os.path.join(output_folder, word_filename)
    document.save(word_path)

# 指定包含 Markdown 文件的文件夹路径
markdown_folder_path = "/path/to/markdown/files"

# 指定输出 Word 文件的文件夹路径
word_output_folder = "/path/to/output/word/files"

# 创建输出文件夹
if not os.path.exists(word_output_folder):
    os.makedirs(word_output_folder)

# 遍历 Markdown 文件并转换为 Word
for filename in os.listdir(markdown_folder_path):
    if filename.endswith(".md"):
        markdown_file = os.path.join(markdown_folder_path, filename)
        convert_markdown_to_word(markdown_file, word_output_folder)
        print(f"已转换文件: {filename}")